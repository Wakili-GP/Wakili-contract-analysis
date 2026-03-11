import modal
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Body, Request
from fastapi.middleware.cors import CORSMiddleware
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
import os
import docx
import shutil
import json
import re
import uuid
import logging
from datetime import datetime, timedelta
from typing import Dict
from dotenv import load_dotenv
load_dotenv()  # ده هيحمل الـ .env تلقائيًا
# LangChain imports
from langchain_community.document_loaders import PyPDFLoader
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Logging setup
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# ────────────────────────────────────────────────
# Modal Configuration
# ────────────────────────────────────────────────

image = (
    modal.Image.debian_slim(python_version="3.12")
    .pip_install(
        "fastapi",
        "uvicorn",
        "slowapi",
        "python-dotenv",
        "python-docx",
        "python-multipart",
        "langchain",
        "langchain-community",
        "langchain-huggingface",
        "langchain-groq",
        "chromadb>=0.4.0,<0.6.0",          # ← هنا الحل: نطاق يتجنب 0.5.4/0.5.5
        "langchain-chroma>=0.1.2",         # ← أحدث شوية عشان يدعم chromadb الحديث
        "sentence-transformers",
        "pypdf",
    )
)

app = modal.App("legal-ai-auditor", image=image)

# Volume لتخزين قواعد Chroma بشكل دائم
chroma_vol = modal.Volume.from_name("chroma-storage", create_if_missing=True)

# الـ FastAPI app الرئيسي
fastapi_app = FastAPI(title="Legal AI Auditor API")

# Rate limiting
limiter = Limiter(key_func=get_remote_address)
fastapi_app.state.limiter = limiter
fastapi_app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

fastapi_app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Cache في الذاكرة (لكل container)
vectorstore_cache: Dict[str, tuple[Chroma, datetime]] = {}


# embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
#embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")  # ← لو حافظتي على local، لازم تحمليه في volume
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
)
llm = ChatGroq(
    temperature=0,
    api_key=os.getenv("GROQ_API_KEY"),
    model_name="llama-3.3-70b-versatile"
)

def read_docx(file_path: str):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text)]

async def create_vectorstore(file: UploadFile, collection_name: str) -> Chroma:
    file_ext = file.filename.split('.')[-1].lower()
    if file_ext not in ["pdf", "docx"]:
        raise HTTPException(400, "برجاء رفع ملف PDF أو DOCX فقط.")

    temp_path = f"/tmp/temp_{uuid.uuid4()}.{file_ext}"
    persist_dir = f"/chroma/{collection_name}"  # داخل الـ volume

    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    try:
        if file_ext == "pdf":
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
        else:
            docs = read_docx(temp_path)

        full_content = " ".join(d.page_content for d in docs)
        if len(full_content.strip()) < 50:
            raise HTTPException(400, "المستند فارغ جداً أو لا يمكن قراءته.")

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=300)
        splits = text_splitter.split_documents(docs)

        vectorstore = Chroma.from_documents(
            documents=splits,
            embedding=embeddings,
            collection_name=collection_name,
            persist_directory=persist_dir
        )

        chroma_vol.commit()  # حفظ التغييرات
        return vectorstore

    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

def cleanup_cache():
    now = datetime.now()
    to_delete = []
    for analysis_id, (vectorstore, expiry) in vectorstore_cache.items():
        if now > expiry:
            to_delete.append(analysis_id)
            try:
                vectorstore.delete_collection()
                logger.info(f"[CLEANUP] Deleted collection for {analysis_id}")
            except Exception as e:
                logger.warning(f"[CLEANUP] Failed: {str(e)}")
    for aid in to_delete:
        del vectorstore_cache[aid]

# ────────────────────────────────────────────────
# Endpoints
# ────────────────────────────────────────────────

@fastapi_app.post("/analyze")
@limiter.limit("10/minute")
async def analyze_contract(request: Request, file: UploadFile = File(...), history: str = Form("[]")):
    cleanup_cache()
    analysis_id = str(uuid.uuid4())
    collection_name = f"contract_{analysis_id.replace('-', '_')[:20]}"

    try:
        vectorstore = await create_vectorstore(file, collection_name)
        expiry = datetime.now() + timedelta(hours=1)
        vectorstore_cache[analysis_id] = (vectorstore, expiry)
        logger.info(f"[ANALYZE] Cached → {analysis_id}")

        retriever = vectorstore.as_retriever(search_kwargs={"k": 10})

        try:
            chat_history = json.loads(history)
            history_text = "\n".join(f"{m['role'].capitalize()}: {m['content']}" for m in chat_history)
        except:
            history_text = ""

        context_docs = retriever.invoke("استخرج نوع العقد والأطراف والمخاطر الرئيسية والالتزامات الجوهرية")
        context = "\n\n".join(d.page_content for d in context_docs)
        if len(context) > 8000:
            context = context[:8000] + "\n[تم تقصير السياق]"

        full_prompt = f"""أنت مستشار قانوني مصري خبير. حلل العقد بناءً على السياق المقدم فقط.
قواعد صارمة:
- لا تستخدم أي معرفة خارج السياق أبدًا.
- حدد نوع العقد بدقة من العنوان أو المقدمة أو العبارات الصريحة فقط.
- إذا لم يكن النوع واضحًا صراحة، اكتب "غير محدد بوضوح في النص".
- رتب المخاطر من الأشد خطورة إلى الأقل.
- صنف كل خطر: "high" أو "medium" أو "low".
- ابدأ الرد مباشرة بالـ {{ بدون أي كلمة أو مسافة قبلها.
- لا تضيف أي نص بعد الـ }} الختامي.
- لا تضع تعليقات أو // داخل الـ JSON.
- ضمن أن الـ JSON صالح 100%.

الصيغة المطلوبة:
{{
  "contract_type": "نوع العقد",
  "summary": "ملخص موجز 60-120 كلمة",
  "risks": [
    {{"level": "high", "description": "الوصف مع السبب", "reference": "البند أو الفقرة"}},
    ...
  ],
  "obligations": {{
    "الطرف_الأول": ["التزام 1", "التزام 2"],
    "الطرف_الثاني": [...]
  }},
  "comparison_table": [
    ["البند", "شروط الطرف الأول", "شروط الطرف الثاني"],
    ["...", "...", "..."]
  ]
}}

السياق المستخرج من العقد:
{context}

تاريخ المحادثة:
{history_text}

الإجابة (ابدأ مباشرة بالـ JSON):"""

        response = llm.invoke(full_prompt)
        content = response.content.strip()

        start = content.find('{')
        end = content.rfind('}') + 1
        if start == -1 or end <= start:
            raise ValueError("LLM did not return valid JSON")

        json_str = content[start:end]
        analysis_dict = json.loads(json_str)

        chroma_vol.commit()  # حفظ إضافي بعد التحليل

        return {
            "analysis_id": analysis_id,
            "filename": file.filename,
            "analysis": analysis_dict,
            "status": "success"
        }

    except Exception as e:
        logger.exception("[ANALYZE] Error")
        raise HTTPException(500, f"خطأ أثناء التحليل: {str(e)}")

@fastapi_app.post("/chat")
@limiter.limit("20/minute")
async def chat_with_contract(request: Request, data: Dict = Body(...)):
    cleanup_cache()
    analysis_id = data.get("analysis_id")
    query = data.get("query")
    history = data.get("history", [])

    if not analysis_id or not query:
        raise HTTPException(400, "analysis_id و query مطلوبين")

    logger.info(f"[CHAT] طلب جديد → analysis_id: {analysis_id} | السؤال: {query[:100]}")

    collection_name = f"contract_{analysis_id.replace('-', '_')[:20]}"
    persist_dir = f"/chroma/{collection_name}"

    try:
        # 1. جرب تحميل من الـ cache أولاً
        if analysis_id in vectorstore_cache:
            vectorstore, expiry = vectorstore_cache[analysis_id]
            logger.info(f"[CHAT] تم العثور على vectorstore في الكاش لـ {analysis_id} (ينتهي في {expiry})")
        else:
            # 2. لو مش موجود في الكاش → حمل من الـ volume (الديسك الدائم)
            logger.info(f"[CHAT] مش موجود في الكاش، جاري تحميل من الـ volume: {collection_name}")
            vectorstore = Chroma(
                collection_name=collection_name,
                embedding_function=embeddings,
                persist_directory=persist_dir
            )
            # أضيفه للكاش عشان المرات الجاية
            expiry = datetime.now() + timedelta(hours=1)
            vectorstore_cache[analysis_id] = (vectorstore, expiry)
            chroma_vol.reload()  # تأكد إن الـ volume محدث
            logger.info(f"[CHAT] تم تحميل vectorstore بنجاح من الـ volume وإضافته للكاش")

        # 3. باقي الكود زي ما كان (retriever + context + prompt + response)
        recent_history = history[-8:] if len(history) >= 8 else history
        history_text = json.dumps(recent_history, ensure_ascii=False, indent=2)

        retriever = vectorstore.as_retriever(search_kwargs={"k": 10})
        logger.info("[CHAT] تم إنشاء retriever بنجاح")

        context_docs = retriever.invoke(query)
        logger.info(f"[CHAT] تم استرجاع {len(context_docs)} جزء/قطعة من العقد")

        context = "\n\n".join(d.page_content for d in context_docs)
        if len(context) > 10000:
            context = context[:10000] + "\n[تم تقصير السياق للحد الأقصى]"

        full_prompt = f"""أنت مستشار قانوني مصري خبير بالقوانين المصرية.

القواعد الصارمة التي يجب اتباعها بدون استثناء:
1. ابدأ ردك مباشرة بالـ JSON بدون أي كلمة أو مسافة أو سطر قبل العلامة {{
2. لا تضع أي نص بعد الـ }} الختامي للـ JSON.
3. لا تضع تعليقات داخل الـ JSON (مثل // أو /* */).
4. ضمن أن الـ JSON صالح نحويًا 100%.

ترتيب الأولوية في الإجابة:
- أولاً: ابحث في "السياق المستخرج من العقد" وابحث عن إجابة واضحة أو مباشرة.
  → إذا وجدت → أجب بناءً عليها فقط واذكر المرجع إن وجد (مثل: حسب البند ٥.٢).
- ثانيًا: إذا لم تجد إجابة واضحة في السياق → استخدم معرفتك بالقانون المصري (مثل قانون العمل رقم ١٢ لسنة ٢٠٠٣، أو القانون المدني، أو غيره) لتقديم إجابة منطقية، لكن يجب أن تذكر صراحة في الإجابة: "غير مذكور في العقد، ولكن حسب القانون المصري..."
- ثالثًا: إذا لم يكن هناك أي أساس عقدي أو قانوني واضح → أجب فقط بـ: "لا توجد معلومات كافية في العقد ولا في القانون المعروف لهذا السؤال."

أجب دائمًا بهذا الشكل فقط:
{{
  "answer": "الإجابة الكاملة والدقيقة هنا (قد تكون فقرة أو أكثر)"
}}

السياق المستخرج من العقد:
{context}

تاريخ المحادثة السابقة:
{history_text}

السؤال الحالي: {query}

الإجابة (ابدأ مباشرة بالـ JSON):"""

        logger.info("[CHAT] جاري إرسال الـ prompt للنموذج...")
        response = llm.invoke(full_prompt)
        content = response.content.strip()

        logger.info(f"[CHAT] رد النموذج الخام: {content[:500]}...")

        # استخراج JSON بطريقة أكثر مرونة
        start = content.find('{')
        end = content.rfind('}') + 1
        if start == -1 or end <= start:
            logger.error("[CHAT] ما لقيناش { أو } في الرد")
            return {"status": "error", "message": "لم يتم العثور على JSON في رد النموذج"}

        json_str = content[start:end]
        try:
            resp_dict = json.loads(json_str)
            answer = resp_dict.get("answer", "")
            if not answer.strip():
                answer = "لا توجد معلومات كافية في العقد"
            logger.info("[CHAT] تم استخراج الإجابة بنجاح")
            return {
            "answer": answer,
            "status": "success",
            "history": recent_history + [{"role": "assistant", "content": answer}]  # أضيفي الرد للـ history
            }
        except json.JSONDecodeError as e:
            logger.error(f"[CHAT] خطأ في تحليل JSON: {str(e)}\nالنص المستخرج: {json_str[:300]}...")
            return {"status": "error", "message": f"خطأ في تحليل رد النموذج: {str(e)}"}

    except Exception as e:
        logger.exception(f"[CHAT] خطأ غير متوقع لـ analysis_id {analysis_id}")
        raise HTTPException(500, f"خطأ داخلي في الدردشة: {str(e)}. شوفي logs السيرفر.")
@fastapi_app.get("/")
async def root():
    return {"status": "online", "active_sessions": len(vectorstore_cache)}

# ────────────────────────────────────────────────
# Modal ASGI entrypoint
# ────────────────────────────────────────────────

@app.function(
    image=image,
    volumes={"/chroma": chroma_vol},
    timeout=600,
    secrets=[modal.Secret.from_name("groq-api-secrets")],
    max_containers=5,
    # keep_warm=1,   # اختياري لو عايزة container دايم شغال (بتكلف فلوس)
)
@modal.asgi_app()
def api():
    return fastapi_app